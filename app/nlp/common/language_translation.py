from langdetect import detect
import fitz
import os
"from docx import Document"

def check_document_language(document):
    """
    If the language of a document is not in english or certain phrases are not in english
    then we move to the other functions 
    """
    for text in document.text.values():
        language = detect(text)
        return language


def check_translated_document_available(document):
    
    # Construct the expected translated filename based on some logic (e.g., appending '_translated')
    translated_filename = document.file_name.replace(".pdf","_translated.pdf")
    print(translated_filename)
    # Check if the translated file exists in the upload folder
    translated_file_path = os.path.abspath(os.path.join("uploads", translated_filename))

    print(translated_file_path)
    print(os.path.isfile(translated_file_path))
    if not os.path.isfile(translated_file_path):
        return False
    else:
        return True

def check_certificate_translation(document):
     # Construct the expected translated filename based on some logic (e.g., appending '_translated')
    oath_file_name = 'oath_of_translator.pdf'
    # Check if the translated file exists in the upload folder
    oath_file_path = os.path.join("uploads", oath_file_name)

    if not os.path.exists(oath_file_path):
        return False
    else:
        form_5_text = """
        Form of Oath by Translator  
        (S.C.R., Order VIII Rule 4) 
        IN THE SUPREME COURT OF INDIA 
        In the matter of  , a translator. 
        I,   solemnly affirm and say that I will translate correctly and accurately all documents 
        given to me for translation. 
        Dated this the   day of   20  
        Before me. 
        Registrar 
        """

        doc = fitz.open(oath_file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()

        print(form_5_text in text)
        # Check if the extracted text contains the Form 5 text
        return form_5_text in text
        """WE WOULD HAVE TO CHECK FOR SIGNATURE"""

"""def check_format_translated(document):
    "" 
    Check whether the foramt of translated document matches with the original document
    ""
    original_doc = Document(document.path)
    translated_doc = Document(document.translated_file_path)
    
    # Initialize a dictionary to store format aspects
    format_check = {
        "paragraph_count": True,
        "heading_count": True,
        "bullet_points_count": True,
        "table_count": True
    }
    
    # Check paragraph count
    if len(original_doc.paragraphs) != len(translated_doc.paragraphs):
        format_check["paragraph_count"] = False
    
    # Check heading count (assuming headings are differentiated by style 'Heading 1', 'Heading 2', etc.)
    original_headings = sum(1 for p in original_doc.paragraphs if p.style.name.startswith('Heading'))
    translated_headings = sum(1 for p in translated_doc.paragraphs if p.style.name.startswith('Heading'))
    if original_headings != translated_headings:
        format_check["heading_count"] = False
    
    # Check bullet points count
    original_bullets = sum(1 for p in original_doc.paragraphs if p.style.name.startswith('List Bullet'))
    translated_bullets = sum(1 for p in translated_doc.paragraphs if p.style.name.startswith('List Bullet'))
    if original_bullets != translated_bullets:
        format_check["bullet_points_count"] = False
    
    # Check table count
    if len(original_doc.tables) != len(translated_doc.tables):
        format_check["table_count"] = False
    
    return format_check
"""

def check_consistency_translation():
    """ Check whether translation is correct"""
    pass

